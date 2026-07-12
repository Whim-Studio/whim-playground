"""Build the ~10-page OpsAgent college final-year PROJECT PROPOSAL (.docx)."""
import os
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docgen import (base_document, heading, para, bullet, numbered, figure,
                    table, code_block, page_break, hrule, add_page_number_footer,
                    NAVY, BLUE, SLATE, GREY)

doc = base_document()

# ---------------------------------------------------------------- TITLE PAGE
def title_line(text, size, color, bold=True, after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size); r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after)
    return p

for _ in range(2):
    doc.add_paragraph()
title_line("PROJECT PROPOSAL", 15, GREY, after=4)
title_line("OpsAgent", 30, NAVY, after=2)
title_line("An LLM-Assisted Serverless Incident-Response System for Amazon ECS",
           14, BLUE, bold=True, after=18, italic=True)

hrule(doc)
title_line("A Final-Year Project Proposal submitted in partial fulfilment of the",
           11, SLATE, bold=False, after=1)
title_line("requirements for the degree of Bachelor of Engineering / Technology",
           11, SLATE, bold=False, after=16)

for label in ["Submitted by:  [Student Name]",
              "Roll No.:  [Roll No]",
              "Guide:  [Guide Name]",
              "Department:  [Department]",
              "College:  [College Name]",
              "Academic Year:  [Year]"]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label)
    r.font.size = Pt(12); r.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(6)
add_page_number_footer(doc, "OpsAgent — Project Proposal")
page_break(doc)

# ---------------------------------------------------------------- ABSTRACT
heading(doc, "1. Abstract", 1)
para(doc,
     "Modern cloud services fail in the small hours when no engineer is watching. The gap between an "
     "alarm firing and a human acknowledging it — the Mean Time To Acknowledge (MTTA) — is where most "
     "avoidable outages grow. OpsAgent is a serverless incident-response system that closes this gap for "
     "workloads running on Amazon Elastic Container Service (ECS). When Amazon CloudWatch raises an alarm, "
     "OpsAgent automatically gathers the relevant logs and metrics, asks a Large Language Model (LLM) — "
     "Cohere's Chat API — to produce a structured Site-Reliability-Engineering (SRE) diagnosis, and, only "
     "when it is highly confident, performs a tightly-scoped remediation such as restarting a task or "
     "scaling a service. Every incident is reported to Slack and stored in Amazon DynamoDB, where a React "
     "dashboard makes the full history visible.")
para(doc,
     "This proposal presents the motivation, objectives, literature context, proposed architecture, "
     "methodology, technology stack, expected outcomes, and timeline for building OpsAgent as a final-year "
     "engineering project. The system deliberately favours safety over autonomy: automatic actions are "
     "limited to two reversible operations, are gated behind a strict confidence threshold, and are bounded "
     "so the agent can never scale a service without limit. The result is a realistic, cloud-native "
     "demonstration of how modern generative AI can be embedded responsibly into operational tooling.")
para(doc, "Keywords: Incident Response, AIOps, Large Language Models, Amazon ECS, AWS Lambda, "
          "Serverless, Site Reliability Engineering, Auto-Remediation, DynamoDB, Cohere.",
     italic=True, color=SLATE)

# ---------------------------------------------------------------- INTRODUCTION
heading(doc, "2. Introduction", 1)
heading(doc, "2.1 Background", 2)
para(doc,
     "Containerised microservices deployed on orchestrators such as Amazon ECS have become the default way "
     "to run production workloads. With scale comes operational noise: dozens of services each emit metrics "
     "and logs, and each can raise alarms at any hour. Traditional monitoring stops at notification — it "
     "tells a human that something is wrong but does nothing to diagnose or fix it. On-call engineers are "
     "therefore paged for issues that follow well-understood patterns (a memory leak, an exhausted "
     "connection pool, a retry storm), spending scarce attention on repetitive triage.")
heading(doc, "2.2 Motivation", 2)
para(doc,
     "The recent maturity of Large Language Models makes it feasible to encode the reasoning an SRE performs "
     "during the first minutes of an incident. An LLM can read a short window of logs and metrics and "
     "propose a probable root cause and a corrective action, together with a confidence score. If that "
     "confidence is high enough and the action is safe and reversible, a machine can apply it immediately — "
     "shrinking MTTA from minutes to seconds. If the confidence is low, the case is escalated to a human "
     "with the diagnosis already attached, saving triage time. OpsAgent is built around exactly this "
     "human-in-the-loop-when-needed principle.")
heading(doc, "2.3 Problem Statement", 2)
para(doc,
     "To design and implement a safe, serverless, cost-efficient system that automatically detects, "
     "diagnoses, and — where confidence and safety permit — remediates operational incidents on Amazon "
     "ECS, while always keeping a complete, auditable record and escalating uncertain cases to human "
     "engineers.")

# ---------------------------------------------------------------- OBJECTIVES
heading(doc, "3. Objectives", 1)
numbered(doc, "Ingest CloudWatch alarms via Amazon SNS and normalise both custom and standard alarm "
              "payloads into a single internal alert representation.")
numbered(doc, "Automatically enrich each alert with the most recent logs and Container-Insights metrics "
              "for the affected service.")
numbered(doc, "Use an LLM to produce a structured JSON diagnosis containing root cause, recommended "
              "action, confidence, and reasoning.")
numbered(doc, "Execute only tightly-scoped, reversible remediations (restart a task, scale a service) and "
              "only when confidence strictly exceeds 0.75.")
numbered(doc, "Escalate every low-confidence or non-actionable case to on-call engineers through Slack.")
numbered(doc, "Persist every incident to DynamoDB and surface the full history through a read-only React "
              "dashboard.")
numbered(doc, "Keep the entire system serverless, least-privilege, and inexpensive to operate.")

# ---------------------------------------------------------------- LITERATURE / EXISTING SYSTEMS
heading(doc, "4. Literature Review & Existing Systems", 1)
para(doc,
     "Automated operations (\"AIOps\") is an established field, but most commercial tools stop at anomaly "
     "detection and alert correlation. The table below positions OpsAgent against representative "
     "categories of existing solutions.")
table(doc,
      ["Approach", "What it does", "Limitation OpsAgent addresses"],
      [["Threshold alerting (raw CloudWatch)", "Fires alarms on metric thresholds", "No diagnosis, no action — pure notification"],
       ["Runbook automation (e.g. SSM)", "Runs fixed scripts on triggers", "Rigid; cannot reason about novel log patterns"],
       ["Commercial AIOps suites", "Correlate and cluster alerts", "Costly, opaque, rarely auto-remediate safely"],
       ["Manual on-call triage", "Human reads logs and decides", "Slow, expensive, error-prone at 3 a.m."],
       ["OpsAgent (proposed)", "LLM diagnosis + gated auto-remediation", "Combines reasoning, safety limits, and audit trail"]],
      widths=[1.7, 2.2, 2.4], font=9.5)
para(doc,
     "OpsAgent's contribution is not a new model but a disciplined integration: it treats the LLM as an "
     "advisory component whose output is constrained by an explicit safety policy, rather than as an "
     "autonomous actor. This mirrors current best practice in production AIOps, where generative reasoning "
     "is paired with bounded, reversible action spaces.")

page_break(doc)

# ---------------------------------------------------------------- PROPOSED SYSTEM / ARCHITECTURE
heading(doc, "5. Proposed System Architecture", 1)
para(doc,
     "OpsAgent is composed of two independent serverless paths. The incident-processing path reacts to "
     "alarms and is write-heavy; the dashboard path is read-only. Separating them keeps the remediation "
     "logic isolated from user-facing traffic and lets each scale on its own.")
figure(doc, "fig_architecture.png",
       "Figure 1: OpsAgent system architecture — event-driven incident path (top) and read-only dashboard path (bottom).")
para(doc, "The principal components are:")
bullet(doc, "Incident Lambda (lambda_handler.py + agent.py) — the event entry point and orchestrator.")
bullet(doc, "Tooling layer (tools.py) — wraps CloudWatch Logs, CloudWatch metrics, and ECS actions.")
bullet(doc, "Prompt builder (prompts.py) — assembles the structured SRE prompt sent to the LLM.")
bullet(doc, "Notifier (slack.py) and persistence (dynamo.py) — report and record every incident.")
bullet(doc, "Dashboard API Lambda (api.py) and React/Vite front-end — expose and visualise incident history.")

heading(doc, "6. System Flow", 1)
para(doc,
     "The end-to-end control flow is deterministic and easy to audit. The decision to act is a single, "
     "explicit gate: the LLM's recommended action must be one of the two permitted operations and its "
     "confidence must strictly exceed the threshold of 0.75.")
figure(doc, "fig_flow.png",
       "Figure 2: Incident-handling flow, from alarm to remediation, notification, and persistence.", width=5.2)
para(doc, "Data moves between four actors and one store, summarised in the Data Flow Diagram below.")
figure(doc, "fig_dfd.png",
       "Figure 3: Level-1 Data Flow Diagram of the incident-processing pipeline.")

# ---------------------------------------------------------------- MODULES
heading(doc, "7. Module Design", 1)
para(doc,
     "The backend is organised into small, single-responsibility Python modules that call one another as "
     "shown below. This makes each concern independently testable and keeps the LLM interaction, the AWS "
     "tool calls, and the safety policy cleanly separated.")
figure(doc, "fig_components.png",
       "Figure 4: Backend module and component view showing dependencies between Python modules and the React client.")
table(doc,
      ["Module", "Responsibility"],
      [["lambda_handler.py", "SNS entry point; normalises alarm payloads into an internal alert"],
       ["agent.py", "Orchestrates enrichment, LLM call, confidence gate, and remediation"],
       ["prompts.py", "Builds the structured SRE diagnosis prompt"],
       ["tools.py", "CloudWatch log/metric reads and scoped ECS restart / scale operations"],
       ["slack.py", "Formats and posts incident and failure notifications"],
       ["dynamo.py", "Writes normalised incident records to DynamoDB"],
       ["api.py", "Read-only dashboard API: /incidents, /incidents/{id}, /stats"]],
      widths=[1.7, 4.6], font=10)

# ---------------------------------------------------------------- SAFETY / POLICY
heading(doc, "8. Safety & Remediation Policy", 1)
para(doc, "Safety is a first-class design goal. The remediation policy is intentionally narrow:")
bullet(doc, "Only two actions are automatable: restart (stop one running task so ECS replaces it) and "
            "scale (increase desired count by two).")
bullet(doc, "Scaling is bounded by a maximum desired count of six, so the agent can never scale without limit.")
bullet(doc, "An action runs only if confidence is strictly greater than 0.75; otherwise the case is escalated.")
bullet(doc, "Any exception during handling triggers an explicit failure notification to Slack.")
para(doc, "The expected LLM response contract is a small JSON object, which keeps parsing robust:")
code_block(doc,
           '{\n'
           '  "root_cause": "concise description of what caused this",\n'
           '  "action": "restart | scale | none",\n'
           '  "confidence": 0.0-1.0,\n'
           '  "reasoning": "why this action was chosen"\n'
           '}')

# ---------------------------------------------------------------- METHODOLOGY
heading(doc, "9. Methodology", 1)
para(doc,
     "The project follows an incremental, iterative methodology. Each iteration delivers a working slice — "
     "from a stub that only logs alerts, through metric enrichment and LLM integration, to gated "
     "remediation and finally the dashboard. This de-risks the LLM and AWS integrations early and keeps a "
     "demonstrable system at every milestone.")
numbered(doc, "Requirement analysis and AWS resource design (SNS, ECS, DynamoDB, IAM).")
numbered(doc, "Implement alert ingestion and context enrichment (logs + metrics).")
numbered(doc, "Integrate the LLM and define the JSON diagnosis contract.")
numbered(doc, "Implement the confidence-gated remediation policy with bounded actions.")
numbered(doc, "Add Slack notification, DynamoDB persistence, and the dashboard API.")
numbered(doc, "Build the React dashboard with a demo-data fallback for offline demonstration.")
numbered(doc, "Test, harden IAM to least privilege, and document.")

# ---------------------------------------------------------------- TECH STACK
heading(doc, "10. Technology Stack", 1)
table(doc,
      ["Layer", "Technology"],
      [["Compute", "AWS Lambda (container image, Python 3.11)"],
       ["Eventing", "Amazon SNS + Amazon CloudWatch Alarms"],
       ["Observability", "CloudWatch Logs, CloudWatch / Container Insights metrics"],
       ["Orchestration", "Amazon ECS (restart & scale operations via boto3)"],
       ["Intelligence", "Cohere Chat API (command-a-plus model family)"],
       ["Persistence", "Amazon DynamoDB (id primary key, service-index GSI)"],
       ["Notification", "Slack incoming webhook"],
       ["Frontend", "React 18 + Vite + Tailwind CSS 4"],
       ["Packaging", "Docker, Amazon ECR"]],
      widths=[1.9, 4.4], font=10)

# ---------------------------------------------------------------- OUTCOMES / SCOPE
heading(doc, "11. Expected Outcomes", 1)
bullet(doc, "A deployable serverless pipeline that turns an ECS alarm into a diagnosed, optionally "
            "auto-remediated, and fully recorded incident.")
bullet(doc, "A measurable reduction in Mean Time To Acknowledge for the classes of incident the agent handles.")
bullet(doc, "A React dashboard providing incident history, statistics, and per-incident detail.")
bullet(doc, "A reusable, safety-first pattern for embedding LLMs in operational tooling.")

heading(doc, "12. Scope & Limitations", 1)
para(doc,
     "OpsAgent is a prototype. It intentionally excludes infrastructure-as-code, a CI pipeline, alert "
     "deduplication, and a dead-letter queue. A successful ECS API call is treated as resolved without "
     "waiting for downstream health recovery, and incident processing is synchronous. These boundaries are "
     "documented so the system is evaluated honestly and so future work is clearly framed.")

heading(doc, "13. Proposed Timeline", 1)
table(doc,
      ["Phase", "Weeks", "Deliverable"],
      [["Analysis & design", "1-2", "Architecture, AWS resource plan, IAM design"],
       ["Ingestion & enrichment", "3-4", "Alert normalisation + logs/metrics retrieval"],
       ["LLM integration", "5-6", "Prompt, diagnosis parsing, confidence gate"],
       ["Remediation & persistence", "7-8", "Scoped ECS actions, Slack, DynamoDB"],
       ["Dashboard", "9-10", "React UI + dashboard API"],
       ["Testing & documentation", "11-12", "Hardening, evaluation, final report"]],
      widths=[2.1, 1.0, 3.2], font=10)

heading(doc, "14. Conclusion", 1)
para(doc,
     "OpsAgent demonstrates that generative AI can be embedded into cloud operations responsibly. By "
     "pairing LLM reasoning with a strict, bounded, reversible action policy and a complete audit trail, "
     "the project delivers a realistic AIOps prototype that is safe to demonstrate and instructive to "
     "study. It shows a practical path from raw alarm to intelligent, accountable response.")

heading(doc, "15. References", 1)
refs = [
    "Amazon Web Services, \"AWS Lambda Developer Guide,\" AWS Documentation, 2024.",
    "Amazon Web Services, \"Amazon ECS Developer Guide,\" AWS Documentation, 2024.",
    "Amazon Web Services, \"Amazon CloudWatch User Guide,\" AWS Documentation, 2024.",
    "Amazon Web Services, \"Amazon DynamoDB Developer Guide,\" AWS Documentation, 2024.",
    "Cohere, \"Chat API Reference,\" Cohere Documentation, 2024.",
    "B. Beyer et al., Site Reliability Engineering: How Google Runs Production Systems, O'Reilly, 2016.",
    "OpsAgent source repository, https://github.com/shubhamaher8/OpsAgent, 2024.",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"[{i}] {r}")
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(3)

out = os.path.join(os.path.dirname(__file__), "OpsAgent_Project_Proposal.docx")
doc.save(out)
print("saved", out)
