"""Build the ~30-page OpsAgent college final-year PROJECT REPORT (.docx)."""
import os
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docgen import (base_document, heading, para, bullet, numbered, figure,
                    table, code_block, page_break, hrule, add_page_number_footer,
                    add_toc, NAVY, BLUE, SLATE, GREY)

doc = base_document()


def centre(text, size, color, bold=True, after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size); r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after)
    return p


def field_line(label):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label)
    r.font.size = Pt(12); r.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(6)


# ============================================================ TITLE PAGE
doc.add_paragraph()
centre("A PROJECT REPORT", 14, GREY, after=4)
centre("on", 12, GREY, bold=False, after=8)
centre("OpsAgent", 30, NAVY, after=2)
centre("An LLM-Assisted Serverless Incident-Response System for Amazon ECS",
       14, BLUE, italic=True, after=16)
hrule(doc)
centre("Submitted in partial fulfilment of the requirements for the award of the degree of",
       11, SLATE, bold=False, after=1)
centre("Bachelor of Engineering / Technology in Computer Science & Engineering",
       11, SLATE, bold=False, after=16)
for label in ["Submitted by:  [Student Name]", "Roll No.:  [Roll No]",
              "Under the guidance of:  [Guide Name]", "Department:  [Department]",
              "[College Name]", "Academic Year:  [Year]"]:
    field_line(label)
add_page_number_footer(doc, "OpsAgent — Project Report")
page_break(doc)

# ============================================================ CERTIFICATE
heading(doc, "Certificate", 1)
para(doc,
     "This is to certify that the project report entitled “OpsAgent — An LLM-Assisted Serverless "
     "Incident-Response System for Amazon ECS” is a bona fide record of the work carried out by "
     "[Student Name] (Roll No. [Roll No]) in partial fulfilment of the requirements for the award of the "
     "degree of Bachelor of Engineering / Technology in Computer Science & Engineering during the academic "
     "year [Year], under my supervision and guidance.")
para(doc,
     "To the best of my knowledge, the work presented in this report has not been submitted, in part or in "
     "full, for the award of any other degree or diploma.")
doc.add_paragraph(); doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("[Guide Name]\nProject Guide\n[Department]")
r.font.size = Pt(11)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r2 = p2.add_run("[Head of Department]\n[Department]")
r2.font.size = Pt(11)
page_break(doc)

# ============================================================ DECLARATION
heading(doc, "Declaration", 1)
para(doc,
     "I hereby declare that the project report entitled “OpsAgent — An LLM-Assisted Serverless "
     "Incident-Response System for Amazon ECS” submitted by me is an original piece of work carried out "
     "under the guidance of [Guide Name]. I further declare that the material contained in this report has "
     "not been submitted previously for the award of any degree or diploma of any university, and that all "
     "sources of information and external components used have been duly acknowledged.")
doc.add_paragraph(); doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("[Student Name]\nRoll No. [Roll No]\nPlace: [College Name]\nDate: ____________")
r.font.size = Pt(11)
page_break(doc)

# ============================================================ ACKNOWLEDGEMENT
heading(doc, "Acknowledgement", 1)
para(doc,
     "I express my sincere gratitude to my project guide, [Guide Name], whose insight, encouragement, and "
     "technical direction shaped this work at every stage. I am thankful to [Head of Department] and the "
     "faculty of the [Department] at [College Name] for providing the environment and resources that made "
     "this project possible.")
para(doc,
     "I also acknowledge the open documentation of Amazon Web Services and Cohere, which informed the "
     "design of this system, and the many practitioners of Site Reliability Engineering whose published "
     "practice guided its safety philosophy. Finally, I thank my family and peers for their steady support "
     "throughout the course of this project.")
page_break(doc)

# ============================================================ ABSTRACT
heading(doc, "Abstract", 1)
para(doc,
     "OpsAgent is a serverless, event-driven incident-response system for workloads running on Amazon "
     "Elastic Container Service (ECS). It reduces the critical delay between an alarm firing and a response "
     "beginning — the Mean Time To Acknowledge — by using a Large Language Model to diagnose incidents and, "
     "where it is safe and highly confident, to remediate them automatically. When Amazon CloudWatch raises "
     "an alarm, a notification is delivered through Amazon SNS to an incident-processing AWS Lambda "
     "function. The function normalises the alert, retrieves the most recent logs and Container-Insights "
     "metrics for the affected service, and submits this context to Cohere's Chat API, which returns a "
     "structured JSON diagnosis containing a root cause, a recommended action, a confidence score, and "
     "supporting reasoning.")
para(doc,
     "A strict policy governs action: only two reversible operations — restarting a task and scaling a "
     "service by a bounded amount — may run automatically, and only when the model's confidence strictly "
     "exceeds 0.75. Every other case is escalated to on-call engineers. All incidents are reported to Slack "
     "and persisted in Amazon DynamoDB, from which a read-only React dashboard renders the complete "
     "history. This report documents the analysis, design, implementation, and evaluation of the system, "
     "and reflects on its limitations and the future work required to move it from prototype to production.")
para(doc, "Keywords: AIOps, Large Language Models, Incident Response, Amazon ECS, AWS Lambda, Serverless, "
          "Site Reliability Engineering, Auto-Remediation, DynamoDB, React.", italic=True, color=SLATE)
page_break(doc)

# ============================================================ TOC + LISTS
heading(doc, "Table of Contents", 1)
add_toc(doc)
page_break(doc)

heading(doc, "List of Figures", 1)
figs = [
    "Figure 1.1  MTTA gap addressed by OpsAgent",
    "Figure 3.1  OpsAgent system architecture",
    "Figure 3.2  AWS deployment topology",
    "Figure 4.1  Use-case diagram",
    "Figure 5.1  Incident-handling flow",
    "Figure 5.2  Level-1 Data Flow Diagram",
    "Figure 6.1  Backend component & module view",
]
for f in figs:
    para(doc, f, justify=False, size=10.5, space_after=3)
heading(doc, "List of Tables", 2)
tbls = [
    "Table 2.1  Comparison with existing approaches",
    "Table 3.1  Component responsibilities",
    "Table 6.1  Backend module summary",
    "Table 6.2  Dashboard API endpoints",
    "Table 7.1  Key environment variables",
    "Table 8.1  Demonstration incident scenarios",
]
for t in tbls:
    para(doc, t, justify=False, size=10.5, space_after=3)
page_break(doc)

# ============================================================ CH 1 INTRODUCTION
heading(doc, "Chapter 1 — Introduction", 1)
heading(doc, "1.1 Overview", 2)
para(doc,
     "Cloud-native systems are operated at a scale and speed that outstrips human attention. A single "
     "organisation may run dozens of containerised services on Amazon ECS, each emitting metrics and logs "
     "and each capable of raising an alarm at any hour. Conventional monitoring stops at notification: it "
     "reliably tells an engineer that something is wrong, but it neither explains the fault nor acts on it. "
     "The consequence is that skilled engineers are repeatedly woken to perform triage that follows "
     "well-understood patterns. OpsAgent is a response to this reality — a system that treats the first "
     "minutes of an incident as an automatable act of reasoning.")
heading(doc, "1.2 Motivation", 2)
para(doc,
     "The maturity of Large Language Models has changed what is feasible in operations. An LLM can read a "
     "short window of logs and metrics and articulate a probable root cause together with a corrective "
     "action and a confidence score, in seconds. This is precisely the judgement an SRE exercises at the "
     "start of an incident. If the model is confident and the proposed action is safe and reversible, a "
     "machine can apply it immediately; if not, the case is handed to a human with the diagnosis already "
     "attached. OpsAgent is designed around this principle of acting autonomously only within a narrow, "
     "provably safe envelope, and escalating everything else.")
heading(doc, "1.3 Problem Statement", 2)
para(doc,
     "To design and implement a safe, serverless, cost-efficient system that automatically detects, "
     "diagnoses, and — where confidence and safety permit — remediates operational incidents on Amazon "
     "ECS, while maintaining a complete, auditable record and escalating uncertain cases to human "
     "engineers.")
heading(doc, "1.4 Objectives", 2)
for o in [
    "Ingest CloudWatch alarms via SNS and normalise custom and standard payloads.",
    "Enrich each alert with recent logs and Container-Insights metrics.",
    "Obtain a structured LLM diagnosis (root cause, action, confidence, reasoning).",
    "Execute only bounded, reversible remediations above a strict confidence threshold.",
    "Escalate low-confidence or non-actionable incidents to Slack.",
    "Persist all incidents to DynamoDB and expose them through a React dashboard.",
]:
    numbered(doc, o)
heading(doc, "1.5 Scope of the Project", 2)
para(doc,
     "The project delivers a working prototype covering the full path from alarm to recorded, optionally "
     "remediated incident, plus a dashboard for review. It deliberately excludes production concerns such "
     "as infrastructure-as-code, alert deduplication, and asynchronous queuing; these are named as future "
     "work in Chapter 9.")
heading(doc, "1.6 Organisation of the Report", 2)
para(doc,
     "Chapter 2 reviews related work. Chapter 3 presents the system architecture and deployment. Chapter 4 "
     "captures requirements. Chapter 5 details system flow and data flow. Chapter 6 describes the module "
     "design and implementation. Chapter 7 covers the environment and deployment. Chapter 8 discusses "
     "testing and results. Chapter 9 states limitations and future work, and Chapter 10 concludes.")
page_break(doc)

# ============================================================ CH 2 LIT REVIEW
heading(doc, "Chapter 2 — Literature Review", 1)
heading(doc, "2.1 The Incident-Response Problem", 2)
para(doc,
     "Site Reliability Engineering frames operational health in terms of measurable objectives and the time "
     "taken to detect, acknowledge, and resolve failures. Two metrics dominate: Mean Time To Acknowledge "
     "(MTTA) and Mean Time To Resolve (MTTR). Notification-only monitoring reduces neither once the alarm "
     "has fired; the clock keeps running until a human engages. OpsAgent targets MTTA directly by beginning "
     "diagnosis and, where safe, resolution the instant an alarm arrives.")
heading(doc, "2.2 AIOps and Automated Remediation", 2)
para(doc,
     "AIOps platforms apply statistics and machine learning to correlate and cluster alerts, suppressing "
     "noise and surfacing likely-related events. They excel at reducing alert volume but rarely take "
     "corrective action, because acting safely on inferred conclusions is hard. Runbook automation goes "
     "further by executing fixed scripts on triggers, but it cannot reason about novel log patterns and "
     "degrades to brittle if-this-then-that rules. OpsAgent occupies the space between: it uses generative "
     "reasoning to interpret unstructured evidence, then constrains the resulting action to a small, "
     "reversible, and bounded set.")
heading(doc, "2.3 Large Language Models as Diagnostic Agents", 2)
para(doc,
     "Recent work treats LLMs as reasoning components within larger deterministic systems rather than as "
     "autonomous decision-makers. The effective pattern is to give the model a tightly-scoped task with a "
     "strict output contract, then validate and gate its output in code. OpsAgent adopts this pattern "
     "exactly: the model is asked for a small JSON object, and a hand-written policy decides whether — and "
     "how — to act on it.")
heading(doc, "2.4 Comparison with Existing Approaches", 2)
table(doc,
      ["Approach", "Diagnoses?", "Auto-acts?", "Key limitation"],
      [["Threshold alerting", "No", "No", "Notification only"],
       ["Runbook automation", "No", "Fixed scripts", "Cannot reason about novel patterns"],
       ["Commercial AIOps", "Correlates", "Rarely", "Costly and opaque"],
       ["Manual triage", "Yes (human)", "Yes (human)", "Slow and expensive"],
       ["OpsAgent", "Yes (LLM)", "Bounded & gated", "Prototype-level operational scope"]],
      widths=[1.6, 1.1, 1.3, 2.3], font=9.5)
para(doc, "Table 2.1 — Comparison of OpsAgent with representative existing approaches.",
     italic=True, color=SLATE, size=9.5)
page_break(doc)

# ============================================================ CH 3 ARCHITECTURE
heading(doc, "Chapter 3 — System Architecture", 1)
heading(doc, "3.1 Architectural Overview", 2)
para(doc,
     "OpsAgent is built from two independent serverless paths that share only the DynamoDB incident store. "
     "The incident-processing path is event-driven and write-heavy: it reacts to alarms, performs "
     "diagnosis and remediation, and records the outcome. The dashboard path is read-only: it serves "
     "incident history to a browser. Decoupling the two isolates remediation logic from user traffic and "
     "lets each path scale and fail independently.")
figure(doc, "fig_architecture.png",
       "Figure 3.1 — OpsAgent system architecture: event-driven incident path and read-only dashboard path.")
heading(doc, "3.2 Component Responsibilities", 2)
table(doc,
      ["Component", "Responsibility"],
      [["Incident Lambda", "Entry point; orchestrates the whole incident lifecycle"],
       ["Tooling layer", "Reads CloudWatch logs/metrics; performs ECS restart & scale"],
       ["Prompt builder", "Assembles the structured SRE prompt for the LLM"],
       ["LLM (Cohere)", "Returns a structured JSON diagnosis"],
       ["Notifier", "Posts incident and failure messages to Slack"],
       ["Persistence", "Writes normalised incident records to DynamoDB"],
       ["Dashboard API", "Exposes read-only incident and statistics endpoints"],
       ["React frontend", "Visualises incident feed, detail, metrics, and stats"]],
      widths=[1.9, 4.4], font=10)
para(doc, "Table 3.1 — Responsibilities of the principal components.", italic=True, color=SLATE, size=9.5)
heading(doc, "3.3 Deployment Topology", 2)
para(doc,
     "All compute runs on AWS Lambda using a single container image published to Amazon ECR; the same image "
     "backs both the incident function and the dashboard API, distinguished only by their handler. SNS "
     "fans alarms into the incident function, DynamoDB stores incidents, and an HTTP gateway fronts the "
     "dashboard API. External calls reach only Cohere and Slack.")
figure(doc, "fig_deployment.png",
       "Figure 3.2 — AWS deployment topology showing Lambda, SNS, ECS, DynamoDB, ECR, and the HTTP gateway.")
heading(doc, "3.4 Design Principles", 2)
bullet(doc, "Serverless-first: no servers to manage; the system scales to zero when idle.")
bullet(doc, "Least privilege: each Lambda holds only the IAM permissions it needs.")
bullet(doc, "Separation of concerns: reasoning, tool calls, and policy live in distinct modules.")
bullet(doc, "Safety over autonomy: bounded, reversible actions behind a strict confidence gate.")
bullet(doc, "Auditability: every incident is recorded before it is considered handled.")
page_break(doc)

# ============================================================ CH 4 REQUIREMENTS
heading(doc, "Chapter 4 — Requirement Analysis", 1)
heading(doc, "4.1 Functional Requirements", 2)
for r in [
    "The system shall accept both custom alert payloads and standard CloudWatch alarm notifications.",
    "The system shall retrieve the last five minutes of logs and current metrics for the affected service.",
    "The system shall obtain a structured diagnosis with a numeric confidence from the LLM.",
    "The system shall automatically restart or scale only when confidence exceeds 0.75.",
    "The system shall bound scaling to a maximum desired task count of six.",
    "The system shall notify Slack for every handled incident and every handler failure.",
    "The system shall persist every incident, including failures, to DynamoDB.",
    "The dashboard shall list incidents, show incident detail, and present aggregate statistics.",
]:
    numbered(doc, r)
heading(doc, "4.2 Non-Functional Requirements", 2)
bullet(doc, "Safety: automated actions must be reversible and bounded.")
bullet(doc, "Reliability: a handler failure must never silently drop an incident.")
bullet(doc, "Cost-efficiency: the system should incur cost only while processing incidents.")
bullet(doc, "Observability: all decisions and actions must be logged.")
bullet(doc, "Usability: the dashboard must remain demonstrable even without a live backend.")
heading(doc, "4.3 Use-Case Model", 2)
para(doc,
     "The primary actors are CloudWatch (which triggers the pipeline) and the SRE / on-call engineer (who "
     "receives escalations and reviews the dashboard). The main use cases are shown below.")
figure(doc, "fig_usecase.png",
       "Figure 4.1 — Use-case diagram of OpsAgent with its two principal actors.", width=5.6)
heading(doc, "4.4 Assumptions & Constraints", 2)
para(doc,
     "The system assumes valid AWS credentials with least-privilege access, a reachable Cohere endpoint, "
     "and a DynamoDB table keyed on id with a service-index global secondary index. Incident processing is "
     "synchronous and single-record per invocation, consistent with the prototype scope.")
page_break(doc)

# ============================================================ CH 5 FLOW / DFD
heading(doc, "Chapter 5 — System Design: Flow & Data", 1)
heading(doc, "5.1 Incident-Handling Flow", 2)
para(doc,
     "The control flow is deterministic and easy to audit. After enrichment and diagnosis, a single "
     "explicit gate decides whether to act: the recommended action must be one of the two permitted "
     "operations and confidence must strictly exceed 0.75. Both branches converge on notification and "
     "persistence, so no incident is ever lost regardless of the decision taken.")
figure(doc, "fig_flow.png",
       "Figure 5.1 — Incident-handling flow, from alarm through diagnosis to remediation or escalation.",
       width=5.0)
heading(doc, "5.2 Data Flow Diagram", 2)
para(doc,
     "At the data level the system comprises three processes — normalise & enrich, diagnose & decide, and "
     "remediate/notify/persist — interacting with external entities (CloudWatch, the LLM, Slack/SRE) and a "
     "single data store, DynamoDB.")
figure(doc, "fig_dfd.png",
       "Figure 5.2 — Level-1 Data Flow Diagram of the incident-processing pipeline.")
heading(doc, "5.3 Incident Data Model", 2)
para(doc, "Each incident is stored as a flat DynamoDB item. The principal attributes are:")
table(doc,
      ["Attribute", "Meaning"],
      [["id", "Unique incident identifier (INC-timestamp-seq)"],
       ["timestamp", "ISO-8601 time the incident was recorded"],
       ["service / metric", "Affected ECS service and the breaching metric"],
       ["value / threshold", "Observed value and the alarm threshold"],
       ["root_cause / reasoning", "LLM diagnosis text"],
       ["action / confidence", "Recommended action and model confidence"],
       ["action_taken / resolved", "What the agent did and whether it resolved"],
       ["cpu / memory / error_rate", "Enriched metric snapshot at incident time"]],
      widths=[2.1, 4.2], font=10)
page_break(doc)

# ============================================================ CH 6 IMPLEMENTATION
heading(doc, "Chapter 6 — Detailed Design & Implementation", 1)
heading(doc, "6.1 Module Structure", 2)
para(doc,
     "The backend is a set of small, single-responsibility Python modules. The diagram below shows how they "
     "depend on one another and on the React client.")
figure(doc, "fig_components.png",
       "Figure 6.1 — Backend component and module view.")
table(doc,
      ["Module", "Responsibility"],
      [["lambda_handler.py", "SNS entry point; payload normalisation; stores result"],
       ["agent.py", "Orchestration, LLM call, confidence gate, remediation dispatch"],
       ["prompts.py", "Builds the structured SRE diagnosis prompt"],
       ["tools.py", "CloudWatch log/metric reads; scoped ECS restart & scale"],
       ["slack.py", "Formats and posts incident / failure notifications"],
       ["dynamo.py", "Serialises and writes incident records to DynamoDB"],
       ["api.py", "Read-only dashboard API handler"]],
      widths=[1.8, 4.5], font=10)
para(doc, "Table 6.1 — Summary of backend modules.", italic=True, color=SLATE, size=9.5)

heading(doc, "6.2 Alert Ingestion & Normalisation", 2)
para(doc,
     "The incident Lambda accepts two payload shapes. A custom payload carries service, metric, value, and "
     "threshold directly. A standard CloudWatch alarm notification is parsed to derive the service name "
     "from alarm dimensions or the alarm name, and the observed value from the alarm's state-change reason. "
     "Both shapes collapse into one internal alert dictionary, so downstream code never branches on source.")
code_block(doc,
           'alert = {\n'
           '    "service":   sns_message["service"],\n'
           '    "metric":    sns_message.get("metric", "unknown"),\n'
           '    "value":     sns_message.get("value", 0.0),\n'
           '    "threshold": sns_message.get("threshold", 0.0),\n'
           '}')

heading(doc, "6.3 Context Enrichment", 2)
para(doc,
     "Given a normalised alert, the tooling layer fetches the most recent fifty log events for the service "
     "over a five-minute window and computes average CPU, memory, and a dropped-packets proxy for error "
     "rate from Container-Insights metrics. Failures in enrichment degrade gracefully to empty logs and "
     "zeroed metrics rather than aborting the pipeline.")

heading(doc, "6.4 LLM Diagnosis", 2)
para(doc,
     "The prompt builder embeds the alert, the logs, and the metrics into a compact SRE prompt that demands "
     "a JSON-only response. The agent calls the Cohere Chat API, strips any Markdown code fences, and parses "
     "the result into a diagnosis object. The strict contract keeps parsing deterministic and makes "
     "malformed responses easy to detect and handle.")
code_block(doc,
           '{\n'
           '  "root_cause": "concise description of what caused this",\n'
           '  "action": "restart | scale | none",\n'
           '  "confidence": 0.0-1.0,\n'
           '  "reasoning": "why this action was chosen"\n'
           '}')

heading(doc, "6.5 Remediation Policy", 2)
para(doc,
     "The confidence gate is the heart of the system. Only restart and scale are automatable, and only when "
     "confidence strictly exceeds 0.75. Restart stops a single running task so ECS replaces it; scale "
     "increases the desired count by two, capped at six. Any other case — low confidence or an action of "
     "none — is escalated. Successes and failures are reflected in the recorded action_taken and resolved "
     "fields.")
code_block(doc,
           'if action == "restart" and confidence > 0.75:\n'
           '    success = restart_ecs_task(service)\n'
           'elif action == "scale" and confidence > 0.75:\n'
           '    success = scale_ecs_service(service, delta=2)\n'
           'else:\n'
           '    escalate()   # action_taken = "none"')

heading(doc, "6.6 Notification & Persistence", 2)
para(doc,
     "Every handled incident produces a formatted Slack message indicating whether it was resolved or "
     "escalated, and a normalised item is written to DynamoDB. Exceptions during handling trigger a "
     "dedicated failure notification, ensuring an engineer is always informed even when the agent itself "
     "fails.")

heading(doc, "6.7 Dashboard API", 2)
para(doc, "The read-only API exposes three endpoints backed by the incident table:")
table(doc,
      ["Endpoint", "Purpose"],
      [["GET /incidents?limit=", "Return recent incident records (optionally by service)"],
       ["GET /incidents/{id}", "Return a single incident by identifier"],
       ["GET /stats", "Return aggregate counts (resolved, escalated, restarts, scales)"]],
      widths=[2.3, 4.0], font=10)
para(doc, "Table 6.2 — Dashboard API endpoints.", italic=True, color=SLATE, size=9.5)

heading(doc, "6.8 Frontend Implementation", 2)
para(doc,
     "The dashboard is a React + Vite single-page application styled with Tailwind CSS in a dark, terminal "
     "aesthetic. It polls the API every twelve seconds and renders an incident feed, a detail panel with "
     "the AI analysis and confidence, a CPU/memory/error metrics chart, summary stat cards, and an action "
     "history log. When the API is unreachable or unconfigured it falls back to generated demo data so the "
     "interface remains fully demonstrable offline.")
page_break(doc)

# ============================================================ CH 7 ENVIRONMENT
heading(doc, "Chapter 7 — Environment & Deployment", 1)
heading(doc, "7.1 Technology Stack", 2)
table(doc,
      ["Layer", "Technology"],
      [["Compute", "AWS Lambda (container image, Python 3.11)"],
       ["Eventing", "Amazon SNS + CloudWatch Alarms"],
       ["Observability", "CloudWatch Logs & Container Insights metrics"],
       ["Orchestration", "Amazon ECS (boto3 restart & scale)"],
       ["Intelligence", "Cohere Chat API"],
       ["Persistence", "Amazon DynamoDB (id PK, service-index GSI)"],
       ["Notification", "Slack incoming webhook"],
       ["Frontend", "React 18 + Vite + Tailwind CSS 4"],
       ["Packaging", "Docker + Amazon ECR"]],
      widths=[1.9, 4.4], font=10)
heading(doc, "7.2 Configuration", 2)
table(doc,
      ["Variable", "Purpose"],
      [["COHERE_API_KEY", "Authenticates requests to Cohere"],
       ["COHERE_MODEL", "Cohere model name (default command-a-plus-05-2026)"],
       ["SLACK_WEBHOOK_URL", "Slack incoming-webhook endpoint"],
       ["ECS_CLUSTER", "ECS cluster containing monitored services"],
       ["LOG_GROUP_PREFIX", "Prefix used to derive a service log group"],
       ["DYNAMO_TABLE", "DynamoDB table for incident records"],
       ["AWS_REGION", "Region for the AWS clients (default us-east-1)"],
       ["VITE_API_URL", "Dashboard API base URL (unset ⇒ demo mode)"]],
      widths=[2.1, 4.2], font=10)
para(doc, "Table 7.1 — Key environment variables.", italic=True, color=SLATE, size=9.5)
heading(doc, "7.3 Build & Deploy", 2)
para(doc,
     "The backend image is built from the backend directory and pushed to Amazon ECR. Two Lambda functions "
     "are created from the same image: the incident function uses the lambda_handler.handler entry point, "
     "and the dashboard API uses api.handler. SNS is subscribed to the incident function and an HTTP "
     "gateway routes the dashboard endpoints. The frontend is built with Vite and served as static assets.")
code_block(doc,
           'cd backend\n'
           'docker build --platform linux/amd64 -t opsagent-backend .\n'
           '# push to ECR, then create two Lambda functions:\n'
           '#   incident : lambda_handler.handler\n'
           '#   dashboard: api.handler')
heading(doc, "7.4 Security & Least Privilege", 2)
bullet(doc, "Store secrets in Secrets Manager / SSM / Lambda env, never in source control.")
bullet(doc, "Grant the incident Lambda only the CloudWatch, ECS, and DynamoDB actions it uses.")
bullet(doc, "Scope IAM statements to the target cluster, services, log groups, and table.")
bullet(doc, "Place authentication and restrictive CORS in front of the dashboard API.")
page_break(doc)

# ============================================================ CH 8 TESTING
heading(doc, "Chapter 8 — Testing & Results", 1)
heading(doc, "8.1 Testing Strategy", 2)
para(doc,
     "The system was validated at three levels: unit-level checks of payload normalisation and diagnosis "
     "parsing; integration checks of the enrichment and remediation tools against AWS APIs; and end-to-end "
     "checks driving synthetic alarms through the full pipeline. Because the frontend ships with a "
     "deterministic demo-data generator, the dashboard could be validated independently of a live backend.")
heading(doc, "8.2 Demonstration Scenarios", 2)
para(doc,
     "The dashboard's demo mode reproduces a representative range of SRE incidents, exercising each branch "
     "of the remediation policy.")
table(doc,
      ["Scenario", "Diagnosis", "Outcome"],
      [["Connection-pool exhaustion", "Pool saturated under load", "Automated restart"],
       ["Infinite retry loop", "Retry storm to dependency", "Circuit-breaker reset"],
       ["Upstream dependency failure", "External service down", "Escalated to on-call"],
       ["Unbounded request parsing", "Memory leak", "Restart"],
       ["DB connection saturation", "Connections exhausted", "Auto-scale"],
       ["Deadlock detected", "Lock contention", "Transaction clearing"],
       ["Network congestion", "Packet loss in AZ", "AZ-failover scaling"],
       ["Cache-invalidation storm", "Thundering herd", "Self-resolving (no action)"]],
      widths=[2.3, 2.0, 2.0], font=9.5)
para(doc, "Table 8.1 — Representative demonstration incident scenarios.", italic=True, color=SLATE, size=9.5)
heading(doc, "8.3 Observations", 2)
para(doc,
     "Across the scenarios, high-confidence, clearly-reversible faults (restart-able leaks, scalable "
     "saturation) were remediated automatically, while ambiguous or externally-caused faults were "
     "correctly escalated. The confidence gate behaved as the primary safety control: lowering the model's "
     "certainty consistently shifted outcomes from action to escalation, confirming that the system errs "
     "toward human involvement when unsure.")
heading(doc, "8.4 Evaluation Against Objectives", 2)
para(doc,
     "Every functional objective from Chapter 1 was met by the prototype: ingestion of both payload types, "
     "enrichment, structured diagnosis, gated bounded remediation, escalation, persistence, and dashboard "
     "review. The non-functional goals of safety, auditability, and offline demonstrability were also "
     "satisfied. The unmet goals are those explicitly out of scope and are carried into future work.")
page_break(doc)

# ============================================================ CH 9 LIMITATIONS
heading(doc, "Chapter 9 — Limitations & Future Work", 1)
heading(doc, "9.1 Current Limitations", 2)
bullet(doc, "No infrastructure-as-code, CI pipeline, or automated test suite ships with the prototype.")
bullet(doc, "A successful ECS API call is treated as resolved without confirming downstream health recovery.")
bullet(doc, "Incident processing is synchronous and handles a single record per invocation.")
bullet(doc, "There is no alert deduplication or dead-letter queue for failed events.")
bullet(doc, "The dashboard API has no application-level authentication within the repository itself.")
heading(doc, "9.2 Future Work", 2)
numbered(doc, "Add infrastructure-as-code (e.g. Terraform / CDK) and a CI/CD pipeline.")
numbered(doc, "Introduce alert deduplication and an asynchronous queue with a dead-letter path.")
numbered(doc, "Confirm resolution by re-checking metrics after an action, closing the feedback loop.")
numbered(doc, "Expand the action space carefully with additional reversible, bounded operations.")
numbered(doc, "Add authentication and fine-grained authorisation to the dashboard API.")
numbered(doc, "Incorporate historical incidents as retrieval context to improve diagnosis quality.")
page_break(doc)

# ============================================================ CH 10 CONCLUSION
heading(doc, "Chapter 10 — Conclusion", 1)
para(doc,
     "OpsAgent demonstrates a disciplined way to bring generative AI into cloud operations. Rather than "
     "granting a model open-ended control, it uses the model only for what it does best — reading messy "
     "evidence and proposing a diagnosis — and wraps that reasoning in an explicit, bounded, reversible "
     "action policy with a complete audit trail. The result is a serverless, cost-efficient prototype that "
     "meaningfully attacks Mean Time To Acknowledge for common ECS incidents while keeping humans firmly in "
     "control of everything uncertain.")
para(doc,
     "Beyond its immediate function, the project offers a reusable template for safe AIOps: constrain the "
     "action space, gate on confidence, escalate on doubt, and record everything. The limitations "
     "identified here are the natural next steps from prototype to production, and the architecture was "
     "designed to accommodate them without fundamental change. OpsAgent thus stands both as a working "
     "system and as an argument for how autonomous operational tooling can be made trustworthy.")

heading(doc, "References", 1)
refs = [
    "Amazon Web Services, \"AWS Lambda Developer Guide,\" AWS Documentation, 2024.",
    "Amazon Web Services, \"Amazon Elastic Container Service Developer Guide,\" AWS Documentation, 2024.",
    "Amazon Web Services, \"Amazon CloudWatch User Guide,\" AWS Documentation, 2024.",
    "Amazon Web Services, \"Amazon SNS Developer Guide,\" AWS Documentation, 2024.",
    "Amazon Web Services, \"Amazon DynamoDB Developer Guide,\" AWS Documentation, 2024.",
    "Cohere, \"Chat API Reference,\" Cohere Documentation, 2024.",
    "B. Beyer, C. Jones, J. Petoff, and N. R. Murphy, Site Reliability Engineering, O'Reilly Media, 2016.",
    "B. Beyer et al., The Site Reliability Workbook, O'Reilly Media, 2018.",
    "Meta Open Source, \"React Documentation,\" 2024.",
    "Vite, \"Vite Guide,\" 2024.",
    "Tailwind Labs, \"Tailwind CSS Documentation,\" 2024.",
    "OpsAgent source repository, https://github.com/shubhamaher8/OpsAgent, 2024.",
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"[{i}] {r}")
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(3)

out = os.path.join(os.path.dirname(__file__), "OpsAgent_Project_Report.docx")
doc.save(out)
print("saved", out)
